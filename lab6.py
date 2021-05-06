from docx import Document
from docx.shared import Inches
document=Document()
while True:
	sw = input('1.Задание 1\n2.Задание 2\n3.Задание 2(вывести информацию из таблицы)\n')
	if(sw == '1'):
		doc=Document()
		doc.add_heading('Начало работы с документами',0)
		text=doc.add_paragraph("Этот текст раздела дополняется словами")
		text.add_run('полужирным').bold=True
		text.add_run('шрифтом и ')
		text.add_run('курсивом.').italic=True
		tx=('После абзаца можно привести пример добавления текста через переменную, с использованием нового стиля.')
		te=('Например список с цифровыми или точечными меркерами. ')
		doc.add_paragraph(tx,style='List Number')
		doc.add_paragraph(te,style='List Number')
		doc.add_paragraph('Далее приведем пример того же самого, но с точкой')
		doc.add_paragraph(tx,style='List Bullet')
		doc.add_paragraph(te,style='List Bullet')
		doc.add_heading('Заголовок первого уровня, например для таблицы',level=1)
		records=(
			(3,'101','Компьюьтер'),
			(7,'422','Сканер'),
			(4,'631','Монитор')
		)
		table=doc.add_table(rows=1,cols=3,style='Light Grid')
		hdr_cells=table.rows[0].cells
		hdr_cells[0].text='№ п/п'
		hdr_cells[1].text='Код товара'
		hdr_cells[2].text='Наименование'
		for qty,id, desc in records:
			row_cells=table.add_row().cells
			row_cells[0].text=str(qty)
			row_cells[1].text=id
			row_cells[2].text=desc

		doc.add_paragraph()
		place=('C:\\Users\\Юрий\\Downloads\\wed\\JSD7W337wyw.jpg')
		doc.add_picture(place,width=Inches(3))
		doc.add_page_break()
		doc.save('DEMO1.docx')
	elif sw == '2':
		title=document.add_heading('',0)
		title.add_run('Группа 230781. ').bold=True
		title.add_run('Кривец Юрий Алексеевич.').italic=True
		document.add_heading('Статистика игроков НБА',level=1)
		txt = [
			   ['Бредли Бил'], 
			   ['Вашингтон Уизардс'], 
			   ['защитник'], 
			   ['31.2'], 
			   ['42'], 
			   ['Джоэл Эмбиид'], 
			   ['Филадельфия Сиксерс'], 
			   ['центровой'], 
			   ['29.9'], 
			   ['33'],
			   ['Стефен Карри'],
			   ['Голден Стэйт Уорриорз'],
			   ['разыгрывающий'],
			   ['29.7'], 
			   ['43'],
			   ['Дэмьен Лиллард'], 
			   ['Портленд Трэйл Блэйзерс'],
			   ['разыгрывающий'],
			   ['29.0'], 
			   ['49'], 
			   ['Яннис Адетокунбо'], 
			   ['Милуоки Бакс'], 
			   ['форвард'], 
			   ['28.8'], 
			   ['45'], 
			   ['Лука Дончич'], 
			   ['Даллас Маверикс'],
			   ['разыгрывающий'], 
			   ['28.5'],
			   ['46'], 
			   ['Кевин Дюрант'],
			   ['Бруклин Нетс'], 
			   ['легкий форвард'],
			   ['28.4'],
			   ['20']]
		text=document.add_paragraph("Пояснение к таблице:\n1 столбец - имя игрока\n2 столбец - название команды\n3 столбец - игровая позиция\n4 столбец - Среднее количество очков за игру\n5 стобец - Количество сыгранных матчей\n")
		table = document.add_table(rows = 7, cols = 5, style = 'Table Grid')
		# k = 0
		# for i in range(7):
		# 	for j in range(5):
		# 		cell = table.cell(i, j)
		# 		cell.text = str(''.join(txt[k]))
		# 		k+=1

		document.save('НБА.docx')
		document.add_paragraph()
		place=('C:\\Users\\Юрий\\Downloads\\wed\\1.jpg')
		with open(place, 'rb') as f:
			a = bytearray(f.read())

		# table.add_picture(place,width=Inches(3))
		cell = table.cell(0, 0)
		cell.text = str(int.from_bytes(a, 'little'))
		document.add_page_break()
		document.save('НБА.docx')
	elif sw=='3':
		table = document.tables[0]
		for row in table.rows:
			s = []
			for cell in row.cells:
				s.append(cell.text)
				print(cell.text)
			document.add_paragraph(' '.join(s), style='List Number')
		document.save('НБА.docx')	

		